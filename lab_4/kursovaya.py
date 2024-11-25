import sys
import os
from PyQt5.QtWidgets import (QApplication, QMainWindow, QLabel, QLineEdit,
                             QPushButton, QComboBox, QWidget, QFileDialog)

from PyQt5 import QtGui
from PIL import Image, ImageDraw, ImageFont
import pandas as pd
from docx import Document
from docx.shared import Pt


class Window(QMainWindow):
    def __init__(self):
        """
        Инициализация главного окна приложения.

        Аргументы:
        None

        Возвращает:
        None
        """
        super(Window, self).__init__()
        self.setWindowTitle("Автозаполнение грамот для ЦВО Творчество")
        self.setGeometry(200, 200, 800, 500)
        self.setWindowIcon(QtGui.QIcon('logo.svg'))
        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)

        self.initUI()

    def initUI(self):
        """
        Инициализация пользовательского интерфейса главного окна.

        Аргументы:
        None

        Возвращает:
        None
        """
        self.text = QLabel("Выберите, какой тип грамот надо создать", self)
        self.text.move(100, 100)
        self.text.adjustSize()

        self.btn = QPushButton("Пользовательские", self)
        self.btn.setFixedWidth(150)
        self.btn.move(100, 150)
        self.btn.clicked.connect(self.selecting_documents_for_the_user)

        self.btn2 = QPushButton("Городские", self)
        self.btn2.setFixedWidth(150)
        self.btn2.move(400, 150)
        self.btn2.clicked.connect(self.selecting_documents_for_the_city)

    def selecting_documents_for_the_user(self):
        """
        Метод для отображения интерфейса выбора файлов и папки
        пользователем.

        Атрибуты:
        self.text : QLabel - текстовое поле для вывода информации
        пользователю.
        self.btn4 : QPushButton - кнопка для выбора PNG файла.
        self.btn5 : QPushButton - кнопка для выбора Excel файла.
        self.btn6 : QPushButton - кнопка для выбора папки для сохранения.
        self.btn7 : QPushButton - кнопка для продолжения к следующему шагу.

        Возвращает:
        None
        """
        self.text.hide()
        self.btn.hide()
        self.btn2.hide()

        self.text21 = QLabel("Выберите PNG картинку для основы, Excel файл для"
                             " входных данных и папку для сохранения.", self)
        self.text21.move(100, 70)
        self.text21.adjustSize()
        self.text21.show()

        self.btn4 = QPushButton("Выбрать PNG файл", self)
        self.btn4.setFixedWidth(400)
        self.btn4.move(100, 100)
        self.btn4.clicked.connect(self.load_png_file)
        self.btn4.show()

        self.btn5 = QPushButton("Выбрать Excel файл", self)
        self.btn5.setFixedWidth(400)
        self.btn5.move(100, 130)
        self.btn5.clicked.connect(self.load_excel_file)
        self.btn5.show()

        self.btn6 = QPushButton("Выбрать папку для сохранения", self)
        self.btn6.setFixedWidth(400)
        self.btn6.move(100, 160)
        self.btn6.clicked.connect(self.save_folder)
        self.btn6.show()

        self.btn7 = QPushButton("Далее", self)
        self.btn7.setFixedWidth(400)
        self.btn7.move(100, 300)
        self.btn7.clicked.connect(self.users_diploms)
        self.btn7.show()

    def load_png_file(self):
        """
        Метод для загрузки PNG файла через диалог выбора файла.

        Атрибуты:
        self.png_file : str - путь к выбранному PNG файлу.
        self.text22 : QLabel - метка, отображающая путь к выбранному PNG файлу.
        Возвращает:
        None
        """
        self.png_file, _ = QFileDialog.getOpenFileName(self, "Выберите PNG "
                                                             "файл", "",
                                                       "png Files (*.png)")
        self.text22 = QLabel(f"Выбран PNG файл: {self.png_file}", self)
        self.text22.move(100, 200)
        self.text22.adjustSize()
        self.text22.show()

    def load_excel_file(self):
        """
        Метод для загрузки Excel файла через диалог выбора файла.

        Атрибуты:
        self.excel_file : str - путь к выбранному Excel файлу.
        self.text23 : QLabel - метка, отображающая путь к выбранному Excel
        файлу.

        Возвращает:
        None
        """
        self.excel_file, _ = QFileDialog.getOpenFileName(self, "Выберите Excel"
                                                               " файл", "",
                                                         "Excel Files (*.xlsx)"
                                                         )
        self.text23 = QLabel(f"Выбран Excel файл: {self.excel_file}", self)
        self.text23.move(100, 230)
        self.text23.adjustSize()
        self.text23.show()

    def save_folder(self):
        """
        Метод для выбора папки для сохранения через диалог выбора папки.

        Атрибуты:
        self.save_folder : str - путь к выбранной папке для сохранения.
        self.text24 : QLabel - метка, отображающая путь к выбранной папке.

        Возвращает:
        None
        """
        self.save_folder = QFileDialog.getExistingDirectory(self,
                                                            "Выберите папку "
                                                            "для сохранения")
        self.text24 = QLabel(f"Выбрана папка: {self.save_folder}", self)
        self.text24.move(100, 260)
        self.text24.adjustSize()
        self.text24.show()

    def users_diploms(self):
        """
        Метод для отображения интерфейса ввода координат и шрифтов для
        генерации дипломов.

        Элементы интерфейса:
        - QLabel для инструкций по вводу координат;
        - QLineEdit для ввода значений координат (X и Y);
        - QComboBox для выбора шрифта и его размера;
        - QLineEdit для ввода ФИО ребенка и педагога;
        - Eлементы для ввода названия образовательного учреждения.

        Возвращает:
        None
        """
        self.text21.hide()
        self.text22.hide()
        self.text23.hide()
        self.text24.hide()

        self.btn4.hide()
        self.btn5.hide()
        self.btn6.hide()
        self.btn7.hide()

        self.text11 = QLabel(
            "Вам необходимо задать X и Y-координаты для расположения текста.",
            self)
        self.text11.move(100, 70)
        self.text11.adjustSize()
        self.text11.show()

        self.text12 = QLabel("Вводить данные необходимо в пикселях.", self)
        self.text12.move(100, 100)
        self.text12.adjustSize()
        self.text12.show()

        self.text13 = QLabel("Ввод для Х:", self)
        self.text13.move(100, 135)
        self.text13.adjustSize()
        self.text13.show()

        self.num1_input = QLineEdit(self)
        self.num1_input.setFixedWidth(50)
        self.num1_input.move(250, 130)
        self.num1_input.show()

        self.text14 = QLabel("Введите данные для Y:", self)
        self.text14.move(100, 200)
        self.text14.adjustSize()
        self.text14.show()

        self.text15 = QLabel("Выберите шрифт:", self)
        self.text15.move(350, 200)
        self.text15.adjustSize()
        self.text15.show()

        self.text110 = QLabel("Введите размер шрифта:", self)
        self.text110.move(600, 200)
        self.text110.adjustSize()
        self.text110.show()

        self.text16 = QLabel("Для ФИО ребенка:", self)
        self.text16.move(100, 235)
        self.text16.adjustSize()
        self.text16.show()

        self.num2_input = QLineEdit(self)
        self.num2_input.setFixedWidth(50)
        self.num2_input.move(250, 230)
        self.num2_input.show()

        self.font_combo = QComboBox(self)
        self.fonts = self.list_fonts()
        self.font_combo.addItems([os.path.basename(f) for f in self.fonts])
        self.font_combo.move(350, 230)
        self.font_combo.resize(200, 30)
        self.font_combo.show()

        self.size_combo = QComboBox(self)
        self.size_combo.addItems(
            ["6", "8", "10", "12", "14", "16", "18", "20", "22", "24", "26",
             "28", "30", "32", "34", "36", "40", "50", "52", "54", "56", "58",
             "60"])
        self.size_combo.move(600, 230)
        self.size_combo.resize(100, 30)
        self.size_combo.show()

        self.text17 = QLabel("Для места:", self)
        self.text17.move(100, 270)
        self.text17.adjustSize()
        self.text17.show()

        self.num3_input = QLineEdit(self)
        self.num3_input.setFixedWidth(50)
        self.num3_input.move(250, 265)
        self.num3_input.show()

        self.font_combo2 = QComboBox(self)
        self.fonts2 = self.list_fonts()
        self.font_combo2.addItems([os.path.basename(f) for f in self.fonts])
        self.font_combo2.move(350, 265)
        self.font_combo2.resize(200, 30)
        self.font_combo2.show()

        self.size_combo2 = QComboBox(self)
        self.size_combo2.addItems(["6", "8", "10", "12", "14", "16", "18",
                                   "20", "22", "24", "26", "28", "30", "32",
                                   "34", "36", "40", "50", "52", "54", "56",
                                   "58", "60"])
        self.size_combo2.move(600, 265)
        self.size_combo2.resize(100, 30)
        self.size_combo2.show()

        self.text18 = QLabel("Для ФИО педагога:", self)
        self.text18.move(100, 305)
        self.text18.adjustSize()
        self.text18.show()

        self.num4_input = QLineEdit(self)
        self.num4_input.setFixedWidth(50)
        self.num4_input.move(250, 300)
        self.num4_input.show()

        self.font_combo3 = QComboBox(self)
        self.fonts3 = self.list_fonts()
        self.font_combo3.addItems([os.path.basename(f) for f in self.fonts])
        self.font_combo3.move(350, 300)
        self.font_combo3.resize(200, 30)
        self.font_combo3.show()

        self.size_combo3 = QComboBox(self)
        self.size_combo3.addItems(
            ["6", "8", "10", "12", "14", "16", "18", "20", "22", "24", "26",
             "28", "30", "32", "34", "36", "40", "50", "52", "54", "56", "58",
             "60"])
        self.size_combo3.move(600, 300)
        self.size_combo3.resize(100, 30)
        self.size_combo3.show()

        self.text19 = QLabel("Для названия ОУ:", self)
        self.text19.move(100, 340)
        self.text19.adjustSize()
        self.text19.show()

        self.num5_input = QLineEdit(self)
        self.num5_input.setFixedWidth(50)
        self.num5_input.move(250, 335)
        self.num5_input.show()

        self.font_combo4 = QComboBox(self)
        self.fonts4 = self.list_fonts()
        self.font_combo4.addItems([os.path.basename(f) for f in self.fonts])
        self.font_combo4.move(350, 335)
        self.font_combo4.resize(200, 30)
        self.font_combo4.show()

        self.size_combo4 = QComboBox(self)
        self.size_combo4.addItems(
            ["6", "8", "10", "12", "14", "16", "18", "20", "22", "24", "26",
             "28", "30", "32", "34", "36", "40", "50", "52", "54", "56", "58",
             "60"])
        self.size_combo4.move(600, 335)
        self.size_combo4.resize(100, 30)
        self.size_combo4.show()

        self.btn3 = QPushButton("Сгенерировать", self)
        self.btn3.setFixedWidth(200)
        self.btn3.move(100, 375)
        self.btn3.clicked.connect(self.creating_diplomas_for_users)
        self.btn3.show()

    def list_fonts(self):
        """
        Метод для получения списка доступных шрифтов на системе.

        Возвращает:
        Список полных путей к файлам шрифтов.
        """
        font_dirs = [
            "C:\\Windows\\Fonts",  # Windows
            "/usr/share/fonts",  # Linux
            "/usr/local/share/fonts",
            "~/.fonts",
            "/Library/Fonts",  # macOS
            "~/Library/Fonts"
        ]
        fonts = []
        for font_dir in font_dirs:
            expanded_dir = os.path.expanduser(font_dir)
            if os.path.exists(expanded_dir):
                for root, dirs, files in os.walk(expanded_dir):
                    for file in files:
                        if file.lower().endswith(".ttf"):
                            fonts.append(os.path.join(root, file))
        return fonts

    def split_text(self, text):
        """
        Метод для раздела заданного текста на строки, у которых длина не
        превышает 30 символов.

        Параметры:
        text (str): Текст, который необходимо разбить на строки.

        Возвращает:
        list: Список строк, каждая из которых не превышает 30 символов.
        """
        words = text.split(' ')
        lines = []
        current_line = ''

        for word in words:
            if len(current_line) + len(word) + 1 > 30:
                lines.append(current_line)
                current_line = word
            else:
                if current_line:
                    current_line += ' ' + word
                else:
                    current_line = word
        if current_line:
            lines.append(current_line)

        return lines

    def creating_diplomas_for_users(self):
        """
        Метод создает дипломы для пользователей на основе введенных данных и
        шаблона изображения.

        Возвращает:
        None: Метод не возвращает значения, результат сохраняется в виде файлов
        изображений.
        """
        x = self.num1_input.text()
        kid = self.num2_input.text()
        place = self.num3_input.text()
        teacher = self.num4_input.text()
        school = self.num5_input.text()
        font_name = self.font_combo.currentText()
        font_name2 = self.font_combo2.currentText()
        font_name3 = self.font_combo3.currentText()
        font_name4 = self.font_combo4.currentText()
        size1 = self.size_combo.currentText()
        size2 = self.size_combo2.currentText()
        size3 = self.size_combo3.currentText()
        size4 = self.size_combo4.currentText()

        x = float(x)
        kid = float(kid)
        place = float(place)
        teacher = float(teacher)
        school = float(school)

        folder = self.save_folder
        image_path = self.png_file
        data = pd.read_excel(self.excel_file)

        font_path = os.path.join("C:\\Windows\\Fonts", font_name)
        try:
            font = ImageFont.truetype(font_path, size=int(size1))
        except IOError:
            print(f"Не удалось загрузить шрифт: {font_path}")
            return

        font_path2 = os.path.join("C:\\Windows\\Fonts", font_name2)
        try:
            font2 = ImageFont.truetype(font_path2, size=int(size2))
        except IOError:
            print(f"Не удалось загрузить шрифт: {font_path2}")
            return

        font_path3 = os.path.join("C:\\Windows\\Fonts", font_name3)
        try:
            font3 = ImageFont.truetype(font_path3, size=int(size3))
        except IOError:
            print(f"Не удалось загрузить шрифт: {font_path3}")
            return

        font_path4 = os.path.join("C:\\Windows\\Fonts", font_name4)
        try:
            font4 = ImageFont.truetype(font_path4, size=int(size4))
        except IOError:
            print(f"Не удалось загрузить шрифт: {font_path4}")
            return

        for index, row in data.iterrows():
            image = Image.open(image_path)
            draw = ImageDraw.Draw(image)

            author_name = row['Фамилия, имя автора конкурсной работы'].replace(
                "\n", " ")
            place_name = row['Место'].replace("\n", " ")
            teacher_name = row['Ф.И.О. педагога (руководителя)'].replace("\n",
                                                                         " ")
            school_name = row[
                'Полное наименование образовательного учреждения'].replace(
                "\n", " ")

            author_lines = self.split_text(author_name)
            place_lines = self.split_text(place_name)
            teacher_lines = self.split_text(teacher_name)
            school_lines = self.split_text(school_name)

            y = kid
            for line in author_lines:
                _, _, _, height = draw.textbbox((0, 0), line, font=font)
                draw.text((x, y), line, anchor="ms", font=font,
                          fill="black")
                y += height

            y = place
            for line in place_lines:
                _, _, _, height = draw.textbbox((0, 0), line, font=font2)
                draw.text((x, y), line, anchor="ms", font=font2,
                          fill="black")
                y += height

            y = teacher
            for line in teacher_lines:
                _, _, _, height = draw.textbbox((0, 0), line, font=font3)
                draw.text((x, y), line, anchor="ms", font=font3,
                          fill="black")
                y += height

            y = school
            for line in school_lines:
                _, _, _, height = draw.textbbox((0, 0), line, font=font4)
                draw.text((x, y), line, anchor="ms", font=font4,
                          fill="black")
                y += height

            filename = f"{school_name}_{teacher_name}_{author_name}.png"

            # Заменяем недопустимые символы в имени файла
            invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*',
                             '\\', '\n', '(', ')', '№', '«', '»', ',', '-']
            for char in invalid_chars:
                filename = filename.replace(char, '')

            if len(filename) > 200:
                filename = filename[:200]
            if not filename.lower().endswith('.png'):
                filename = f"{filename}.png"
            full_path = os.path.join(folder, filename)

            image.save(full_path)

        self.text6 = QLabel("Грамоты созданы", self)
        self.text6.move(100, 430)
        self.text6.adjustSize()
        self.text6.show()

    def selecting_documents_for_the_city(self):
        """
        Метод отображает интерфейс для выбора документов и папки для
        сохранения.

        Атрибуты:
        self.text21 : QLabel - текстовое поле для вывода информации
        пользователю.
        self.btn4 : QPushButton - кнопка для выбора Word файла.
        self.btn5 : QPushButton - кнопка для выбора Excel файла.
        self.btn6 : QPushButton - кнопка для выбора папки для сохранения.
        self.btn7 : QPushButton - кнопка для продолжения к следующему шагу.

        Возвращает:
        None
        """
        self.text.hide()
        self.btn.hide()
        self.btn2.hide()

        self.text21 = QLabel(
            "Выберите Word документ для основы, Excel файл для входных данных "
            "и папку для сохранения.",
            self)
        self.text21.move(100, 70)
        self.text21.adjustSize()
        self.text21.show()

        self.btn4 = QPushButton("Выбрать Word файл", self)
        self.btn4.setFixedWidth(400)
        self.btn4.move(100, 100)
        self.btn4.clicked.connect(self.load_word_file)
        self.btn4.show()

        self.btn5 = QPushButton("Выбрать Excel файл", self)
        self.btn5.setFixedWidth(400)
        self.btn5.move(100, 130)
        self.btn5.clicked.connect(self.load_excel_file)
        self.btn5.show()

        self.btn6 = QPushButton("Выбрать папку для сохранения", self)
        self.btn6.setFixedWidth(400)
        self.btn6.move(100, 160)
        self.btn6.clicked.connect(self.save_folder)
        self.btn6.show()

        self.btn7 = QPushButton("Далее", self)
        self.btn7.setFixedWidth(400)
        self.btn7.move(100, 300)
        self.btn7.clicked.connect(self.creating_diplomas_for_city)
        self.btn7.show()

    def load_word_file(self):
        """
        Загружает Word файл с помощью диалогового окна выбора файла.

        Использует QFileDialog для выбора файла с расширением .docx.
        После выбора отображает путь к выбранному файлу в QLabel.

        Возвращает:
        None
        """
        self.word_file, _ = QFileDialog.getOpenFileName(self,
                                                        "Выберите Word файл",
                                                        "",
                                                        "Word Files (*.docx)")
        self.text24 = QLabel(f"Выбран Word файл: {self.word_file}", self)
        self.text24.move(100, 200)
        self.text24.adjustSize()
        self.text24.show()

    def creating_diplomas_for_city(self):
        """
        Метод создает дипломы для пользователей на основе данных из Excel файла
        и шаблона Word документа.

        Возвращает:
        None: Метод не возвращает значения, результат сохраняется в виде файлов
        Word.
        """
        data = pd.read_excel(self.excel_file)
        folder = self.save_folder

        for index, row in data.iterrows():
            author_name = row['Фамилия, имя автора конкурсной работы'].replace(
                "\n", " ")
            place_name = row['Место'].replace("\n", " ")
            teacher_name = row['Ф.И.О. педагога (руководителя)'].replace("\n",
                                                                         " ")
            school_name = row[
                'Полное наименование образовательного учреждения'].replace(
                "\n", " ")

            doc = Document(self.word_file)

            p6 = doc.paragraphs[3]
            run_teacher = p6.add_run(author_name)
            run_teacher.font.size = Pt(26)
            run_teacher.font.name = 'Times New Roman'
            run_teacher.font.bold = True

            p6 = doc.paragraphs[5]
            run_teacher = p6.add_run(school_name)
            run_teacher.font.size = Pt(18)
            run_teacher.font.name = 'Times New Roman'

            p6 = doc.paragraphs[6]
            run_teacher = p6.add_run(teacher_name)
            run_teacher.font.size = Pt(18)
            run_teacher.font.name = 'Times New Roman'

            p6 = doc.paragraphs[8]
            run_teacher = p6.add_run(place_name)
            run_teacher.font.size = Pt(26)
            run_teacher.font.name = 'Times New Roman'
            run_teacher.font.bold = True

            filename = f"{school_name}_{teacher_name}_{author_name}.docx"

            # Заменяем недопустимые символы в имени файла
            invalid_chars = ['<', '>', ':', '"', '/', '\\', '|', '?', '*',
                             '\\', '\n', '(', ')', '№', '«', '»', ',', '-']
            for char in invalid_chars:
                filename = filename.replace(char, '')

            if len(filename) > 200:
                filename = filename[:200]
            if not filename.lower().endswith('.docx'):
                filename = f"{filename}.docx"
            full_path = os.path.join(folder, filename)

            doc.save(full_path)

        self.text6 = QLabel("Грамоты созданы", self)
        self.text6.move(100, 430)
        self.text6.adjustSize()
        self.text6.show()


def application():
    """
    Инициализирует и запускает приложение Qt.

    Возвращает:
    None: Функция не возвращает значений, управляет жизнью приложения.
    """
    app = QApplication(sys.argv)
    app.setStyleSheet('QLabel { font: bold }')
    win = Window()
    win.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    application()